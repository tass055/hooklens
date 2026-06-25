import asyncio
from datetime import date
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocumentError(RuntimeError):
    pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1, color: RGBColor | None = None) -> None:
    para = doc.add_heading(text, level=level)
    if color:
        for run in para.runs:
            run.font.color.rgb = color


def _add_cell_text(cell, text: str, bold: bool = False) -> None:
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_left_border_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with a left border to create a callout effect."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(11)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "6366F1")  # Indigo
    pBdr.append(left)
    pPr.append(pBdr)

    # Light indigo background shading on the paragraph
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF2FF")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

    # Add indentation
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    pPr.append(ind)


def _add_transcript_with_highlight(doc: Document, full_text: str, start: int, end: int) -> None:
    """Render full transcript with the hook segment highlighted in yellow."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]

    before = full_text[:start]
    hook = full_text[start:end]
    after = full_text[end:]

    if before:
        run = para.add_run(before)
        run.font.size = Pt(10)

    if hook:
        run = para.add_run(hook)
        run.font.size = Pt(10)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    if after:
        run = para.add_run(after)
        run.font.size = Pt(10)


def _build_doc(
    transcript_data: dict[str, Any],
    hook_identification: dict[str, Any],
    hook_scores: dict[str, Any],
    output_path: str,
    video_name: str,
) -> None:
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # -------------------------------------------------------------------------
    # 1. Title page
    # -------------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("HookLens Analysis Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)  # Deep indigo

    doc.add_paragraph()  # Spacer

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(str(video_name)[:120])
    name_run.font.size = Pt(13)
    name_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 2. Executive Summary
    # -------------------------------------------------------------------------
    _add_heading(doc, "Executive Summary", level=1)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    headers = [("Hook Type", hook_identification.get("hook_type", "")),
               ("Estimated Duration", hook_identification.get("hook_duration_estimate", "")),
               ("Overall Score", f"{hook_scores.get('overall_score', 0)} / 10")]
    for i, (label, value) in enumerate(headers):
        _set_cell_shading(table.rows[i].cells[0], "F1F5F9")
        _add_cell_text(table.rows[i].cells[0], label, bold=True)
        _add_cell_text(table.rows[i].cells[1], value)

    doc.add_paragraph()
    rationale_para = doc.add_paragraph()
    rationale_run = rationale_para.add_run(hook_identification.get("rationale", ""))
    rationale_run.italic = True
    rationale_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # -------------------------------------------------------------------------
    # 3. Score Breakdown
    # -------------------------------------------------------------------------
    _add_heading(doc, "Score Breakdown", level=1)

    score_table = doc.add_table(rows=4, cols=3)
    score_table.style = "Table Grid"

    # Header row
    for cell, text in zip(score_table.rows[0].cells, ["Dimension", "Score", "Explanation"]):
        _set_cell_shading(cell, "4F46E5")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    rows_data = [
        ("Overall", str(hook_scores.get("overall_score", 0)), "—"),
        ("Value Proposition", str(hook_scores.get("value_proposition", {}).get("score", 0)),
         hook_scores.get("value_proposition", {}).get("explanation", "")),
        ("Emotional Pull", str(hook_scores.get("emotional_pull", {}).get("score", 0)),
         hook_scores.get("emotional_pull", {}).get("explanation", "")),
    ]
    for i, (dim, score, explanation) in enumerate(rows_data):
        _set_cell_shading(score_table.rows[i + 1].cells[0], "F8FAFC")
        _add_cell_text(score_table.rows[i + 1].cells[0], dim)
        _add_cell_text(score_table.rows[i + 1].cells[1], score, bold=True)
        _add_cell_text(score_table.rows[i + 1].cells[2], explanation)

    # -------------------------------------------------------------------------
    # 4. Strengths
    # -------------------------------------------------------------------------
    _add_heading(doc, "Strengths", level=1, color=RGBColor(0x16, 0xA3, 0x4A))

    for item in hook_scores.get("strengths", []):
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(str(item))
        run.font.size = Pt(11)

    # -------------------------------------------------------------------------
    # 5. Weaknesses
    # -------------------------------------------------------------------------
    _add_heading(doc, "Weaknesses", level=1, color=RGBColor(0xD9, 0x77, 0x06))

    for item in hook_scores.get("weaknesses", []):
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(str(item))
        run.font.size = Pt(11)

    # -------------------------------------------------------------------------
    # 6. Improvement Suggestions
    # -------------------------------------------------------------------------
    _add_heading(doc, "Improvement Suggestions", level=1)

    for item in hook_scores.get("improvement_suggestions", []):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(str(item))
        run.font.size = Pt(11)

    # -------------------------------------------------------------------------
    # 7. Rewritten Hook Example
    # -------------------------------------------------------------------------
    _add_heading(doc, "Rewritten Hook Example", level=1)
    doc.add_paragraph()
    rewritten = hook_scores.get("rewritten_hook_example", "")
    if rewritten:
        _add_left_border_paragraph(doc, rewritten)

    # Alternative Hook Ideas
    alternative_hooks = hook_scores.get("alternative_hooks", [])
    if alternative_hooks:
        doc.add_paragraph()
        _add_heading(doc, "Alternative Hook Ideas", level=2)
        for hook in alternative_hooks:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(str(hook))
            run.font.size = Pt(11)

    # Key Teachings & Viral Quotes
    key_teachings = hook_scores.get("key_teachings", [])
    if key_teachings:
        doc.add_paragraph()
        _add_heading(doc, "Key Teachings & Viral Quotes", level=2)
        for quote in key_teachings:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(f'"{quote.strip()}"')
            run.italic = True
            run.font.size = Pt(11)

    # -------------------------------------------------------------------------
    # 8. Full Transcript Details
    # -------------------------------------------------------------------------
    doc.add_page_break()
    _add_heading(doc, "Transcript Details", level=1)

    # 8.1 Corrected Transcript
    _add_heading(doc, "Corrected Transcript", level=2)
    full_text = transcript_data.get("full_text", "")
    start = hook_identification.get("hook_start_char", 0)
    end = hook_identification.get("hook_end_char", 0)

    legend_para = doc.add_paragraph()
    legend_run = legend_para.add_run("Highlighted text = identified hook segment")
    legend_run.font.size = Pt(9)
    legend_run.font.italic = True
    legend_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    legend_para.add_run("   ")

    _add_transcript_with_highlight(doc, full_text, start, end)

    # 8.2 Original Verbatim Transcript
    original_text = transcript_data.get("original_text")
    if original_text:
        doc.add_paragraph()
        _add_heading(doc, "Original Verbatim Transcript", level=2)
        orig_para = doc.add_paragraph()
        orig_para.style = doc.styles["Normal"]
        run = orig_para.add_run(original_text)
        run.font.size = Pt(10)

    doc.save(output_path)


async def generate_document(
    transcript_data: dict[str, Any],
    hook_identification: dict[str, Any],
    hook_scores: dict[str, Any],
    output_path: str,
    video_name: str,
) -> None:
    try:
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            _build_doc,
            transcript_data,
            hook_identification,
            hook_scores,
            output_path,
            video_name,
        )
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError("Report generation failed. Please try again.") from exc
